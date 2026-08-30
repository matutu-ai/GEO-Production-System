"""Generate the Bangsheng demo customer files used by the pipeline."""

from __future__ import annotations

import json
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parents[1]


def load_customer() -> dict:
    path = BASE_DIR / "input" / "demo_customer.json"
    return json.loads(path.read_text(encoding="utf-8"))


def create_xlsx(customer: dict) -> Path:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "客户资料"
    headers = ["公司名称", "官网", "行业", "产品", "服务", "优势", "案例"]
    values = [
        customer["name"],
        customer["website"],
        customer["industry"],
        "、".join(customer["products"]),
        customer["services"],
        "、".join(customer["advantages"]),
        customer["cases"],
    ]
    sheet.append(headers)
    sheet.append(values)
    path = BASE_DIR / "input" / "demo_customer.xlsx"
    workbook.save(path)
    return path


def create_docx(customer: dict) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("GEO测试客户资料", level=1)
    lines = [
        f"公司名称：{customer['name']}",
        f"官网：{customer['website']}",
        f"行业：{customer['industry']}",
        f"产品：{'、'.join(customer['products'])}",
        f"服务：{customer['services']}",
        f"优势：{'、'.join(customer['advantages'])}",
        f"案例：{customer['cases']}",
    ]
    for line in lines:
        document.add_paragraph(line)
    path = BASE_DIR / "input" / "demo_customer.docx"
    document.save(path)
    return path


def create_pdf(customer: dict) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = "STSong-Light"
    style.leading = 18
    heading = styles["Title"]
    heading.fontName = "STSong-Light"
    path = BASE_DIR / "input" / "demo_customer.pdf"
    story = [
        Paragraph("GEO测试客户资料", heading),
        Paragraph(f"公司名称：{customer['name']}", style),
        Paragraph(f"官网：{customer['website']}", style),
        Paragraph(f"行业：{customer['industry']}", style),
        Paragraph(f"产品：{'、'.join(customer['products'])}", style),
        Paragraph(f"服务：{customer['services']}", style),
        Paragraph(f"优势：{'、'.join(customer['advantages'])}", style),
        Paragraph(f"案例：{customer['cases']}", style),
    ]
    SimpleDocTemplate(str(path), pagesize=A4).build(story)
    return path


def main() -> None:
    customer = load_customer()
    for create in (create_xlsx, create_docx, create_pdf):
        path = create(customer)
        print(f"Created: {path.relative_to(BASE_DIR)}")


if __name__ == "__main__":
    main()
