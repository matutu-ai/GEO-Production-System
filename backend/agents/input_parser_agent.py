"""Input Parser Agent: customer files/URLs -> standardized customer data."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List

BASE_DIR = Path(__file__).resolve().parents[1]

FIELD_LABELS = {
    "公司名称": "name",
    "企业名称": "name",
    "客户名称": "name",
    "官网": "website",
    "官网地址": "website",
    "网址": "website",
    "行业": "industry",
    "所属行业": "industry",
    "产品": "products",
    "主营产品": "products",
    "产品体系": "products",
    "服务": "services",
    "服务内容": "services",
    "优势": "advantages",
    "核心优势": "advantages",
    "案例": "cases",
    "客户案例": "cases",
    "客户": "customers",
    "目标客户": "customers",
}

LIST_FIELDS = {"products", "services", "advantages", "cases", "customers"}


class InputParserAgent:
    task = "input_parsing"

    def run(self, input_path: str) -> Dict[str, Any]:
        try:
            raw_path = str(input_path or "")
            if raw_path.startswith(("http://", "https://")):
                customer_data, raw_information = self._parse_url(raw_path)
            else:
                path = Path(raw_path)
                if not path.is_absolute():
                    path = BASE_DIR / path
                if not path.exists():
                    return self._error(f"input file not found: {path}")
                suffix = path.suffix.lower()
                if suffix == ".xlsx":
                    customer_data, raw_information = self._parse_xlsx(path)
                elif suffix == ".docx":
                    customer_data, raw_information = self._parse_docx(path)
                elif suffix == ".pdf":
                    customer_data, raw_information = self._parse_pdf(path)
                else:
                    return self._error(
                        "unsupported input type, expected .xlsx/.docx/.pdf or URL"
                    )

            result = self._normalize(customer_data, raw_information)
            return {
                "task": self.task,
                "status": "success",
                "confidence": self._confidence(result),
                "result": result,
                "next_action": "company_analysis",
            }
        except Exception as exc:  # pragma: no cover - defensive error boundary
            return self._error(str(exc))

    def _parse_xlsx(self, path: Path) -> tuple:
        from openpyxl import load_workbook

        workbook = load_workbook(path, data_only=True, read_only=True)
        rows = []
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows(values_only=True):
                rows.append(
                    ["" if cell is None else str(cell).strip() for cell in row]
                )
        workbook.close()

        customer_data = {}
        raw_lines = []
        for row_index, row in enumerate(rows):
            if not any(row):
                continue
            raw_lines.append(" | ".join(cell for cell in row if cell))
            for col, cell in enumerate(row):
                field = FIELD_LABELS.get(cell.strip())
                if not field:
                    continue
                for next_row in rows[row_index + 1 :]:
                    if col < len(next_row) and next_row[col]:
                        self._set_value(customer_data, field, next_row[col])
                        break
                if (
                    col + 1 < len(row)
                    and row[col + 1]
                    and FIELD_LABELS.get(row[col + 1].strip()) is None
                ):
                    self._set_value(customer_data, field, row[col + 1])
        return customer_data, "\n".join(raw_lines)

    def _parse_docx(self, path: Path) -> tuple:
        from docx import Document

        document = Document(str(path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(
                    " | ".join(cell.text.strip() for cell in row.cells)
                )
        raw_information = "\n".join(line for line in lines if line)
        return self._extract_structured(raw_information), raw_information

    def _parse_pdf(self, path: Path) -> tuple:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        raw_information = "\n".join(pages).strip()
        return self._extract_structured(raw_information), raw_information

    def _parse_url(self, url: str) -> tuple:
        import requests
        from bs4 import BeautifulSoup

        response = requests.get(
            url,
            timeout=20,
            headers={"User-Agent": "Mozilla/5.0 GEO Production Tool"},
        )
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
            tag.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else ""
        body_lines = [
            re.sub(r"\s+", " ", line).strip()
            for line in soup.get_text("\n", strip=True).splitlines()
            if line.strip()
        ]
        raw_information = "\n".join(body_lines[:200])
        customer_data = self._extract_structured(raw_information)
        customer_data["website"] = url
        if not customer_data.get("name") and title:
            candidate = re.split(r"[_\-|]", title, maxsplit=1)[0].strip()
            if candidate:
                customer_data["name"] = candidate
        return customer_data, raw_information

    def _extract_structured(self, raw_information: str) -> Dict[str, Any]:
        customer_data: Dict[str, Any] = {
            "name": "",
            "website": "",
            "industry": "",
            "products": [],
            "services": [],
            "advantages": [],
            "cases": [],
            "customers": [],
            "raw_information": raw_information.strip(),
        }
        for line in raw_information.splitlines():
            cleaned = line.strip()
            if not cleaned or "：" not in cleaned and ":" not in cleaned:
                continue
            parts = re.split(r"[:：]\s*", cleaned, maxsplit=1)
            if len(parts) < 2:
                continue
            label = parts[0].strip()
            field = FIELD_LABELS.get(label)
            if field:
                self._set_value(customer_data, field, parts[1].strip())
        return customer_data

    def _set_value(
        self,
        customer_data: Dict[str, Any],
        field: str,
        value: str,
    ) -> None:
        if field in LIST_FIELDS:
            items = self._split(value)
            existing = customer_data.get(field, [])
            for item in items:
                if item and item not in existing:
                    existing.append(item)
            customer_data[field] = existing
        else:
            if value and not customer_data.get(field):
                customer_data[field] = value

    def _split(self, value: Any) -> List[str]:
        if value is None:
            return []
        parts = re.split(r"[、,，;；/|\n]+", str(value).strip())
        return [part.strip() for part in parts if part.strip()]

    def _normalize(
        self,
        customer_data: Dict[str, Any],
        raw_information: str,
    ) -> Dict[str, Any]:
        result = {
            "name": customer_data.get("name", ""),
            "website": customer_data.get("website", ""),
            "industry": customer_data.get("industry", ""),
            "products": customer_data.get("products", []) or [],
            "services": customer_data.get("services", []) or [],
            "advantages": customer_data.get("advantages", []) or [],
            "cases": customer_data.get("cases", []) or [],
            "customers": customer_data.get("customers", []) or [],
            "raw_information": raw_information.strip(),
        }
        if not result["name"]:
            first_line = next(
                (
                    line.strip()
                    for line in raw_information.splitlines()
                    if line.strip()
                ),
                "",
            )
            result["name"] = self._first_company_name(first_line)
        return result

    @staticmethod
    def _first_company_name(first_line: str) -> str:
        for marker in ("企业介绍", "企业简介", "公司介绍", "公司简介"):
            if marker in first_line:
                candidate = first_line.split(marker, 1)[0].strip()
                if candidate:
                    return candidate
        return first_line

    def _confidence(self, result: Dict[str, Any]) -> int:
        score = 50
        if result.get("name"):
            score += 15
        if result.get("industry"):
            score += 10
        if result.get("products"):
            score += 15
        if result.get("raw_information"):
            score += 10
        return min(100, score)

    def _error(self, message: str) -> Dict[str, Any]:
        return {
            "task": self.task,
            "status": "error",
            "confidence": 0,
            "result": {
                "name": "",
                "website": "",
                "industry": "",
                "products": [],
                "services": [],
                "advantages": [],
                "cases": [],
                "customers": [],
                "raw_information": "",
            },
            "next_action": "fix_input",
            "message": message,
        }
