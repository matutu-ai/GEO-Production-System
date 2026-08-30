"""Report Agent: enterprise GEO delivery report (docx + pdf)."""

from __future__ import annotations

import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List
from xml.sax.saxutils import escape

from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

BASE_DIR = Path(__file__).resolve().parents[1]
REPORT_VERSION = "GEO Production Tool V1.4"
FONT_NAME = "Microsoft YaHei"


class ReportAgent:
    task = "report_generation"

    def run(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        company = input_data.get("company_profile") or {}
        customer_data = input_data.get("customer_data") or {}
        business = self._result_body(
            input_data.get("business_result") or input_data.get("business_analysis")
        )
        keyword_result = self._result_body(input_data.get("keyword_result"))
        persona_result = self._result_body(input_data.get("persona_result"))
        content_result = self._result_body(input_data.get("content_result"))
        strategy_result = self._result_body(input_data.get("strategy_result"))
        output_dir = Path(input_data.get("output_dir", BASE_DIR / "output"))
        output_dir.mkdir(parents=True, exist_ok=True)

        data = {
            "company": company,
            "customer": customer_data,
            "business_lines": business.get("business_lines", []),
            "keywords": keyword_result.get("keywords", []),
            "personas": persona_result.get("personas", []),
            "directions": content_result.get("content_directions", []),
            "faq_list": content_result.get("faq_list", []),
            "case_list": content_result.get("case_list", []),
            "plan": content_result.get("plan", []),
            "actions": strategy_result.get("priority_actions", []),
            "roadmap": strategy_result.get("30_day_plan", []),
            "strategy_summary": strategy_result.get("summary", ""),
            "report_date": datetime.date.today().isoformat(),
            "company_name": (
                company.get("company_name")
                or customer_data.get("name")
                or "客户企业"
            ),
        }

        try:
            docx_path = self._write_docx(
                output_dir / "GEO客户分析报告.docx", data
            )
            pdf_path = self._write_pdf(
                output_dir / "GEO客户分析报告.pdf", data
            )
        except Exception as exc:
            return self._error(str(exc))

        return {
            "task": self.task,
            "status": "success",
            "confidence": 100,
            "result": {
                "files": [str(docx_path), str(pdf_path)],
            },
            "next_action": "done",
        }

    def _result_body(self, item: Any) -> Dict[str, Any]:
        if isinstance(item, dict) and "result" in item:
            return item.get("result", {}) or {}
        return item or {}

    def _text(self, value: Any) -> str:
        if isinstance(value, list):
            return "、".join(
                str(item) for item in value if str(item).strip()
            )
        return "" if value is None else str(value)

    def _error(self, message: str) -> Dict[str, Any]:
        return {
            "task": self.task,
            "status": "error",
            "confidence": 0,
            "result": {"files": []},
            "next_action": "fix_report_generation",
            "message": message,
        }

    def _write_docx(self, path: Path, data: Dict[str, Any]) -> Path:
        from docx import Document

        template_path = BASE_DIR / "report_templates" / "geo_report_template.docx"
        document = (
            Document(str(template_path))
            if template_path.exists()
            else Document()
        )
        self._fill_cover(document, data)

        self._docx_section(
            document,
            "1 项目概览",
            self._docx_project_overview,
            data,
            page_break=False,
        )
        self._docx_section(
            document, "2 企业分析", self._docx_company_analysis, data
        )
        self._docx_section(
            document,
            "3 业务机会分析",
            self._docx_business_analysis,
            data,
        )
        self._docx_section(
            document,
            "4 GEO关键词策略",
            self._docx_keyword_strategy,
            data,
        )
        self._docx_section(
            document,
            "5 用户画像分析",
            self._docx_persona_analysis,
            data,
        )
        self._docx_section(
            document,
            "6 内容增长计划",
            self._docx_content_plan,
            data,
        )
        self._docx_section(
            document,
            "7 GEO优化执行策略",
            self._docx_strategy,
            data,
        )
        self._docx_section(
            document,
            "8 30天执行路线图",
            self._docx_roadmap,
            data,
        )
        document.save(path)
        return path

    def _fill_cover(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        replacements = {
            "{{REPORT_TITLE}}": "GEO客户分析报告",
            "{{COMPANY_NAME}}": data["company_name"],
            "{{REPORT_DATE}}": data["report_date"],
            "{{REPORT_VERSION}}": REPORT_VERSION,
        }
        for paragraph in document.paragraphs:
            text = paragraph.text
            if not any(token in text for token in replacements):
                continue
            for token, value in replacements.items():
                text = text.replace(token, value)
            if not paragraph.runs:
                continue
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""

    def _docx_section(
        self,
        document: Any,
        title: str,
        builder: Callable[[Any, Dict[str, Any]], None],
        data: Dict[str, Any],
        page_break: bool = True,
    ) -> None:
        if page_break:
            document.add_page_break()
        self._docx_heading(document, title, level=1)
        builder(document, data)

    def _docx_heading(
        self, document: Any, text: str, level: int = 1
    ) -> Any:
        heading = document.add_heading(text, level=level)
        size = 16 if level == 1 else 13
        for run in heading.runs:
            self._set_docx_run_font(run, size=size, bold=True)
        return heading

    def _docx_paragraph(
        self, document: Any, text: str, bold: bool = False
    ) -> Any:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        self._set_docx_run_font(run, bold=bold)
        return paragraph

    def _set_docx_run_font(
        self, run: Any, size: int = 10, bold: bool = False
    ) -> None:
        from docx.oxml.ns import qn
        from docx.shared import Pt

        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT_NAME)

    def _docx_table(
        self,
        document: Any,
        headers: List[str],
        rows: List[List[Any]],
    ) -> Any:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = document.styles["Table Grid"]
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            self._fill_docx_cell(
                header_cells[index], header, bold=True, fill="D9E2F3"
            )
        for row_data in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row_data):
                self._fill_docx_cell(cells[index], self._text(value))
        if not rows:
            cells = table.add_row().cells
            cells[0].text = "暂无数据"
        return table

    def _fill_docx_cell(
        self,
        cell: Any,
        value: Any,
        bold: bool = False,
        fill: str = "",
    ) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        cell.text = self._text(value)
        for run in cell.paragraphs[0].runs:
            self._set_docx_run_font(run, size=10, bold=bold)
        if fill:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)
            tc_pr.append(shading)

    def _docx_project_overview(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        company = data["company"]
        customer = data["customer"]
        self._docx_paragraph(
            document,
            "本报告用于指导客户在 ChatGPT、Claude、Gemini、DeepSeek 等"
            "AI 搜索中的品牌曝光、内容覆盖与询盘转化。",
            bold=False,
        )
        self._docx_table(
            document,
            ["项目信息", "内容"],
            [
                ["企业名称", data["company_name"]],
                ["官网", customer.get("website") or company.get("website") or "未提供"],
                ["行业", company.get("industry", "")],
                ["产品体系", company.get("products", [])],
                ["服务能力", company.get("services", [])],
                ["目标客户", company.get("target_customers", [])],
                ["核心优势", company.get("advantages", [])],
                ["报告版本", REPORT_VERSION],
                ["报告日期", data["report_date"]],
            ],
        )

    def _docx_company_analysis(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        company = data["company"]
        self._docx_table(
            document,
            ["分析维度", "GEO Profile"],
            [
                ["企业定位", company.get("company_positioning", "")],
                ["行业", company.get("industry", "")],
                ["产品体系", company.get("products", [])],
                ["服务能力", company.get("services", [])],
                ["目标客户", company.get("target_customers", [])],
                ["核心优势", company.get("advantages", [])],
                ["客户痛点", company.get("customer_pain_points", [])],
                ["可信证据", company.get("evidence", [])],
            ],
        )

    def _docx_business_analysis(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        lines = data["business_lines"]
        self._docx_paragraph(
            document,
            f"识别到 {len(lines)} 条业务机会，作为关键词、内容和案例策略的业务基础。",
        )
        for index, line in enumerate(lines, start=1):
            name = line.get("business_name") or self._text(
                line.get("products", [])
            )
            self._docx_heading(document, f"{index}. {name}", level=2)
            self._docx_table(
                document,
                ["维度", "内容"],
                [
                    ["产品体系", line.get("products", [])],
                    ["目标客户", line.get("target_customers", [])],
                    ["客户问题", line.get("customer_problems", [])],
                    ["购买意图", line.get("buying_intent", [])],
                    ["关键词方向", line.get("keywords_direction", [])],
                    ["内容方向", line.get("content_direction", [])],
                ],
            )

    def _docx_keyword_strategy(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        keywords = data["keywords"]
        a_count = sum(1 for item in keywords if item.get("priority") == "A")
        self._docx_paragraph(
            document,
            f"共识别 {len(keywords)} 个关键词，其中 A 级高价值词 {a_count} 个，"
            "覆盖品牌、业务、问答、场景和长尾搜索场景。",
        )
        rows = [
            [
                item.get("keyword", ""),
                item.get("type", ""),
                item.get("intent", ""),
                item.get("priority", ""),
                item.get("business_line", ""),
                item.get("customer_type", ""),
                item.get("search_stage", ""),
            ]
            for item in keywords
        ]
        self._docx_table(
            document,
            ["关键词", "类型", "意图", "优先级", "业务线", "客户类型", "搜索阶段"],
            rows,
        )

    def _docx_persona_analysis(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        personas = data["personas"]
        self._docx_paragraph(
            document,
            "围绕企业决策链建立统一画像，用于校准内容主题、案例证据与发布渠道。",
        )
        for persona in personas:
            role = persona.get("role", "目标用户")
            self._docx_heading(document, role, level=2)
            self._docx_table(
                document,
                ["关注点", "痛点", "搜索行为", "决策因素", "内容需求"],
                [
                    [
                        persona.get("focus", []),
                        persona.get("pain_points", []),
                        persona.get("search_behavior", []),
                        persona.get("decision_factors", []),
                        persona.get("content_needs", []),
                    ]
                ],
            )

    def _docx_content_plan(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        self._docx_heading(document, "内容方向", level=2)
        self._docx_table(
            document,
            ["内容方向", "内容主题", "目标关键词", "目标用户", "发布建议"],
            [
                [
                    item.get("direction", ""),
                    item.get("content_topic", ""),
                    item.get("target_keyword", ""),
                    item.get("target_user", ""),
                    item.get("publish_suggestion", ""),
                ]
                for item in data["directions"]
            ],
        )
        self._docx_heading(document, "FAQ 与案例清单", level=2)
        self._docx_table(
            document,
            ["序号", "FAQ问题"],
            [
                [str(index + 1), item]
                for index, item in enumerate(data["faq_list"])
            ],
        )
        self._docx_table(
            document,
            ["序号", "案例主题"],
            [
                [str(index + 1), item]
                for index, item in enumerate(data["case_list"])
            ],
        )
        self._docx_heading(document, "30天内容计划", level=2)
        self._docx_table(
            document,
            ["日期", "内容主题", "目标关键词", "目标用户", "发布建议"],
            [
                [
                    item.get("日期", ""),
                    item.get("内容主题", ""),
                    item.get("目标关键词", ""),
                    item.get("目标用户", ""),
                    item.get("发布建议", ""),
                ]
                for item in data["plan"]
            ],
        )

    def _docx_strategy(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        self._docx_paragraph(
            document,
            data["strategy_summary"] or "暂无策略摘要",
        )
        for priority in ("P1", "P2", "P3"):
            actions = [
                item
                for item in data["actions"]
                if item.get("priority") == priority
            ]
            if not actions:
                continue
            self._docx_heading(document, f"{priority} 行动", level=2)
            self._docx_table(
                document,
                ["执行动作", "原因", "相关关键词", "目标用户", "内容需要", "预计价值"],
                [
                    [
                        item.get("action", ""),
                        item.get("reason", ""),
                        item.get("related_keywords", []),
                        item.get("target_users", []),
                        item.get("content_needed", ""),
                        item.get("expected_value", ""),
                    ]
                    for item in actions
                ],
            )

    def _docx_roadmap(
        self, document: Any, data: Dict[str, Any]
    ) -> None:
        self._docx_paragraph(
            document,
            "以下路线图将策略动作落到周维度，确保30天内完成第一批GEO基础建设。",
        )
        self._docx_table(
            document,
            ["周", "任务"],
            [
                [item.get("week", ""), item.get("tasks", [])]
                for item in data["roadmap"]
            ],
        )

    def _write_pdf(self, path: Path, data: Dict[str, Any]) -> Path:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        base = ParagraphStyle(
            "Chinese",
            fontName="STSong-Light",
            fontSize=10,
            leading=16,
            wordWrap="CJK",
        )
        styles = {
            "title": ParagraphStyle(
                "Title", parent=base, fontSize=26, leading=34,
                alignment=1, spaceAfter=12,
            ),
            "subtitle": ParagraphStyle(
                "Subtitle", parent=base, fontSize=12, leading=18,
                alignment=1, spaceAfter=30,
            ),
            "h1": ParagraphStyle(
                "H1", parent=base, fontSize=18, leading=26, spaceAfter=10,
            ),
            "h2": ParagraphStyle(
                "H2", parent=base, fontSize=13, leading=20,
                spaceBefore=8, spaceAfter=6,
            ),
            "body": ParagraphStyle(
                "Body", parent=base, fontSize=10, leading=16, spaceAfter=6,
            ),
            "cell": ParagraphStyle(
                "Cell", parent=base, fontSize=8.5, leading=12, wordWrap="CJK",
            ),
            "header_cell": ParagraphStyle(
                "HeaderCell",
                parent=base,
                fontSize=8.5,
                leading=12,
                wordWrap="CJK",
                textColor=colors.HexColor("#1F3864"),
            ),
        }
        story: List[Any] = []
        story.append(Paragraph("GEO客户分析报告", styles["title"]))
        story.append(
            Paragraph(
                "企业级 Generative Engine Optimization 客户交付报告",
                styles["subtitle"],
            )
        )
        story.append(
            Paragraph(
                f"交付企业：{escape(data['company_name'])}",
                styles["body"],
            )
        )
        story.append(
            Paragraph(f"报告日期：{escape(data['report_date'])}", styles["body"])
        )
        story.append(
            Paragraph(f"报告版本：{escape(REPORT_VERSION)}", styles["body"])
        )
        story.append(PageBreak())

        self._pdf_project_overview(story, data, styles)
        self._pdf_section(
            story, "2 企业分析", self._pdf_company_analysis, data, styles
        )
        self._pdf_section(
            story,
            "3 业务机会分析",
            self._pdf_business_analysis,
            data,
            styles,
        )
        self._pdf_section(
            story,
            "4 GEO关键词策略",
            self._pdf_keyword_strategy,
            data,
            styles,
        )
        self._pdf_section(
            story,
            "5 用户画像分析",
            self._pdf_persona_analysis,
            data,
            styles,
        )
        self._pdf_section(
            story, "6 内容增长计划", self._pdf_content_plan, data, styles
        )
        self._pdf_section(
            story, "7 GEO优化执行策略", self._pdf_strategy, data, styles
        )
        self._pdf_section(
            story, "8 30天执行路线图", self._pdf_roadmap, data, styles
        )

        document = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
            title=f"{data['company_name']} GEO客户分析报告",
            author=REPORT_VERSION,
        )
        document.build(story)
        return path

    def _pdf_section(
        self,
        story: List[Any],
        title: str,
        builder: Callable[[List[Any], Dict[str, Any], Dict[str, Any]], None],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        story.append(PageBreak())
        story.append(Paragraph(title, styles["h1"]))
        story.append(Spacer(1, 6))
        builder(story, data, styles)

    def _pdf_table(
        self,
        story: List[Any],
        headers: List[str],
        rows: List[List[Any]],
        styles: Dict[str, Any],
        col_widths: List[float],
    ) -> None:
        from reportlab.lib import colors
        from reportlab.platypus import Paragraph, Table, TableStyle

        if not rows:
            story.append(Paragraph("暂无数据", styles["body"]))
            story.append(Spacer(1, 8))
            return
        data = [
            [Paragraph(escape(header), styles["header_cell"]) for header in headers]
        ]
        data.extend(
            [
                [Paragraph(escape(self._text(value)), styles["cell"]) for value in row]
                for row in rows
            ]
        )
        table = Table(data, repeatRows=1, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9AA7B8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 10))

    def _pdf_project_overview(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        company = data["company"]
        customer = data["customer"]
        story.append(Paragraph("1 项目概览", styles["h1"]))
        story.append(
            Paragraph(
                "本报告用于指导客户在 ChatGPT、Claude、Gemini、DeepSeek 等"
                "AI 搜索中的品牌曝光、内容覆盖与询盘转化。",
                styles["body"],
            )
        )
        self._pdf_table(
            story,
            ["项目信息", "内容"],
            [
                ["企业名称", data["company_name"]],
                ["官网", customer.get("website") or company.get("website") or "未提供"],
                ["行业", company.get("industry", "")],
                ["产品体系", company.get("products", [])],
                ["服务能力", company.get("services", [])],
                ["目标客户", company.get("target_customers", [])],
                ["核心优势", company.get("advantages", [])],
                ["报告版本", REPORT_VERSION],
                ["报告日期", data["report_date"]],
            ],
            styles,
            [110, 350],
        )

    def _pdf_company_analysis(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        company = data["company"]
        self._pdf_table(
            story,
            ["分析维度", "GEO Profile"],
            [
                ["企业定位", company.get("company_positioning", "")],
                ["行业", company.get("industry", "")],
                ["产品体系", company.get("products", [])],
                ["服务能力", company.get("services", [])],
                ["目标客户", company.get("target_customers", [])],
                ["核心优势", company.get("advantages", [])],
                ["客户痛点", company.get("customer_pain_points", [])],
                ["可信证据", company.get("evidence", [])],
            ],
            styles,
            [110, 350],
        )

    def _pdf_business_analysis(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        lines = data["business_lines"]
        story.append(
            Paragraph(
                f"识别到 {len(lines)} 条业务机会，作为关键词、内容和案例策略的业务基础。",
                styles["body"],
            )
        )
        for index, line in enumerate(lines, start=1):
            name = line.get("business_name") or self._text(
                line.get("products", [])
            )
            story.append(Paragraph(f"{index}. {escape(name)}", styles["h2"]))
            self._pdf_table(
                story,
                ["维度", "内容"],
                [
                    ["产品体系", line.get("products", [])],
                    ["目标客户", line.get("target_customers", [])],
                    ["客户问题", line.get("customer_problems", [])],
                    ["购买意图", line.get("buying_intent", [])],
                    ["关键词方向", line.get("keywords_direction", [])],
                    ["内容方向", line.get("content_direction", [])],
                ],
                styles,
                [110, 350],
            )

    def _pdf_keyword_strategy(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        keywords = data["keywords"]
        a_count = sum(1 for item in keywords if item.get("priority") == "A")
        story.append(
            Paragraph(
                f"共识别 {len(keywords)} 个关键词，其中 A 级高价值词 {a_count} 个，"
                "覆盖品牌、业务、问答、场景和长尾搜索场景。",
                styles["body"],
            )
        )
        rows = [
            [
                item.get("keyword", ""),
                item.get("type", ""),
                item.get("intent", ""),
                item.get("priority", ""),
                item.get("business_line", ""),
                item.get("customer_type", ""),
                item.get("search_stage", ""),
            ]
            for item in keywords
        ]
        self._pdf_table(
            story,
            ["关键词", "类型", "意图", "优先级", "业务线", "客户类型", "搜索阶段"],
            rows,
            styles,
            [100, 45, 60, 28, 90, 55, 60],
        )

    def _pdf_persona_analysis(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        story.append(
            Paragraph(
                "围绕企业决策链建立统一画像，用于校准内容主题、案例证据与发布渠道。",
                styles["body"],
            )
        )
        for persona in data["personas"]:
            role = persona.get("role", "目标用户")
            story.append(Paragraph(escape(role), styles["h2"]))
            self._pdf_table(
                story,
                ["关注点", "痛点", "搜索行为", "决策因素", "内容需求"],
                [
                    [
                        persona.get("focus", []),
                        persona.get("pain_points", []),
                        persona.get("search_behavior", []),
                        persona.get("decision_factors", []),
                        persona.get("content_needs", []),
                    ]
                ],
                styles,
                [70, 90, 90, 90, 100],
            )

    def _pdf_content_plan(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        story.append(Paragraph("内容方向", styles["h2"]))
        self._pdf_table(
            story,
            ["内容方向", "内容主题", "目标关键词", "目标用户", "发布建议"],
            [
                [
                    item.get("direction", ""),
                    item.get("content_topic", ""),
                    item.get("target_keyword", ""),
                    item.get("target_user", ""),
                    item.get("publish_suggestion", ""),
                ]
                for item in data["directions"]
            ],
            styles,
            [60, 125, 80, 70, 105],
        )
        story.append(Paragraph("FAQ 与案例清单", styles["h2"]))
        self._pdf_table(
            story,
            ["序号", "FAQ问题"],
            [
                [str(index + 1), item]
                for index, item in enumerate(data["faq_list"])
            ],
            styles,
            [40, 420],
        )
        self._pdf_table(
            story,
            ["序号", "案例主题"],
            [
                [str(index + 1), item]
                for index, item in enumerate(data["case_list"])
            ],
            styles,
            [40, 420],
        )
        story.append(Paragraph("30天内容计划", styles["h2"]))
        self._pdf_table(
            story,
            ["日期", "内容主题", "目标关键词", "目标用户", "发布建议"],
            [
                [
                    item.get("日期", ""),
                    item.get("内容主题", ""),
                    item.get("目标关键词", ""),
                    item.get("目标用户", ""),
                    item.get("发布建议", ""),
                ]
                for item in data["plan"]
            ],
            styles,
            [45, 130, 80, 60, 125],
        )

    def _pdf_strategy(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        story.append(
            Paragraph(
                escape(data["strategy_summary"] or "暂无策略摘要"),
                styles["body"],
            )
        )
        for priority in ("P1", "P2", "P3"):
            actions = [
                item
                for item in data["actions"]
                if item.get("priority") == priority
            ]
            if not actions:
                continue
            story.append(Paragraph(f"{priority} 行动", styles["h2"]))
            self._pdf_table(
                story,
                ["执行动作", "原因", "相关关键词", "目标用户", "内容需要", "预计价值"],
                [
                    [
                        item.get("action", ""),
                        item.get("reason", ""),
                        item.get("related_keywords", []),
                        item.get("target_users", []),
                        item.get("content_needed", ""),
                        item.get("expected_value", ""),
                    ]
                    for item in actions
                ],
                styles,
                [90, 80, 70, 60, 70, 70],
            )

    def _pdf_roadmap(
        self,
        story: List[Any],
        data: Dict[str, Any],
        styles: Dict[str, Any],
    ) -> None:
        story.append(
            Paragraph(
                "以下路线图将策略动作落到周维度，确保30天内完成第一批GEO基础建设。",
                styles["body"],
            )
        )
        self._pdf_table(
            story,
            ["周", "任务"],
            [
                [item.get("week", ""), item.get("tasks", [])]
                for item in data["roadmap"]
            ],
            styles,
            [70, 390],
        )
