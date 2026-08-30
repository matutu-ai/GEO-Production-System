"""Generate the GEO delivery report template used by ReportAgent."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(__file__).resolve().parents[1]
FONT_NAME = "Microsoft YaHei"


def _set_style_font(style, size=None) -> None:
    style.font.name = FONT_NAME
    if size is not None:
        style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_run_font(run, size=None, bold=False, color=None) -> None:
    run.font.name = FONT_NAME
    run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def create_template() -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    _set_style_font(document.styles["Normal"], size=11)
    _set_style_font(document.styles["Title"], size=24)
    _set_style_font(document.styles["Heading 1"], size=18)
    _set_style_font(document.styles["Heading 2"], size=14)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run("{{REPORT_TITLE}}")
    _set_run_font(run, size=30, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(60)
    run = subtitle.add_run("企业级 Generative Engine Optimization 客户交付报告")
    _set_run_font(run, size=13)

    for label in (
        "交付企业：{{COMPANY_NAME}}",
        "报告日期：{{REPORT_DATE}}",
        "报告版本：{{REPORT_VERSION}}",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(label)
        _set_run_font(run, size=12)

    document.add_page_break()
    path = BASE_DIR / "report_templates" / "geo_report_template.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def main() -> None:
    path = create_template()
    print(f"Created: {path.relative_to(BASE_DIR)}")


if __name__ == "__main__":
    main()
